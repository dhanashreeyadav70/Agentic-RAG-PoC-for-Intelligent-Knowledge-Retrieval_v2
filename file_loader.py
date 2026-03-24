import os
import pandas as pd
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader, Docx2txtLoader, UnstructuredHTMLLoader
)
from langchain_core.documents import Document
from ingestion import load_json


# def load_file(file_path, filename):

#     ext = os.path.splitext(filename)[1].lower()

#     if ext == ".pdf":
#         return PyPDFLoader(file_path).load()

#     elif ext == ".txt":
#         return TextLoader(file_path).load()

    # elif ext == ".docx":
    #     return Docx2txtLoader(file_path).load()

    # elif ext == ".html":
    #     return UnstructuredHTMLLoader(file_path).load()

    # elif ext == ".json":
    #     return load_json(file_path)

    # # ⭐ FIXED CSV HANDLING
    # elif ext == ".csv":

    #     df = pd.read_csv(file_path)

    #     documents = []

#         # ✅ Dataset-level understanding
#         summary = f"""
# This dataset represents an employee directory.

# Columns: {', '.join(df.columns)}

# This dataset contains structured employee information used for HR, reporting, and organizational analysis.
# """

#         documents.append(Document(
#             page_content=summary,
#             metadata={"source": filename, "type": "summary"}
#         ))

    #     # ✅ Row-level semantic conversion
    #     for _, row in df.iterrows():

    #         row_text = ", ".join([
    #             f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])
    #         ])

    #         documents.append(Document(
    #             page_content=f"Employee record: {row_text}",
    #             metadata={"source": filename, "type": "row"}
    #         ))

    #     return documents

    # else:
    #     raise ValueError(f"Unsupported file type: {ext}")


import os
import json
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from PIL import Image
import pytesseract
import whisper
from moviepy.editor import VideoFileClip


# -------------------------
# Main Loader
# -------------------------
def load_file(file_path, filename):
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return load_pdf(file_path)

    elif ext == ".docx":
        return load_docx(file_path)

    elif ext == ".txt":
        return load_text(file_path)

    elif ext == ".csv":
        return load_csv(file_path)

    elif ext == ".json":
        return load_json(file_path)

    elif ext in [".html", ".htm"]:
        return load_html(file_path)

    elif ext in [".png", ".jpg", ".jpeg"]:
        return load_image(file_path)

    elif ext in [".mp3", ".wav"]:
        return load_audio(file_path)

    elif ext in [".mp4", ".avi", ".mpeg4"]:
        return load_video(file_path)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


# -------------------------
# Document Loaders
# -------------------------
def load_text(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return format_output(text, file_path)


def load_csv(file_path):
    df = pd.read_csv(file_path)
    text = df.to_string()
    return format_output(text, file_path)


def load_json(file_path):
    with open(file_path, "r") as f:
        data = json.load(f)
    text = json.dumps(data, indent=2)
    return format_output(text, file_path)


def load_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return format_output(text, file_path)


def load_pdf(file_path):
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    text = "\n".join([page.extract_text() or "" for page in reader.pages])
    return format_output(text, file_path)


def load_html(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")
    text = soup.get_text()
    return format_output(text, file_path)


# -------------------------
# Image Loader (OCR)
# -------------------------
def load_image(file_path):
    text = pytesseract.image_to_string(Image.open(file_path))
    return format_output(text, file_path)


# -------------------------
# Audio Loader (Whisper)
# -------------------------
def load_audio(file_path):
    model = whisper.load_model("base")
    result = model.transcribe(file_path)
    return format_output(result["text"], file_path)


# -------------------------
# Video Loader
# -------------------------
def load_video(file_path):
    video = VideoFileClip(file_path)

    audio_path = file_path + ".wav"
    video.audio.write_audiofile(audio_path)

    model = whisper.load_model("base")
    result = model.transcribe(audio_path)

    return format_output(result["text"], file_path)


# -------------------------
# Helper
# -------------------------
def format_output(text, source):
    return [{
        "page_content": text,
        "metadata": {"source": source}
    }]
